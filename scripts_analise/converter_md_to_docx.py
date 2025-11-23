#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para converter APRESENTACAO_PDF.md para .docx
Usa python-docx para conversão básica de Markdown
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
from pathlib import Path

# Caminhos
SCRIPT_DIR = Path(__file__).parent
ROOT_DIR = SCRIPT_DIR.parent
MARKDOWN_FILE = ROOT_DIR / "APRESENTACAO_PDF.md"
OUTPUT_FILE = ROOT_DIR / "APRESENTACAO_PDF.docx"

def criar_documento():
    """Cria documento Word com estilos configurados"""
    doc = Document()
    
    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    return doc

def adicionar_titulo(doc, texto, nivel=1):
    """Adiciona título ao documento"""
    heading = doc.add_heading(texto, level=nivel)
    return heading

def adicionar_paragrafo(doc, texto, negrito=False, italico=False, codigo=False):
    """Adiciona parágrafo com formatação"""
    p = doc.add_paragraph()
    
    if codigo:
        run = p.add_run(texto)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
    else:
        # Processar markdown inline (negrito, itálico, código)
        partes = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', texto)
        for parte in partes:
            if parte.startswith('**') and parte.endswith('**'):
                run = p.add_run(parte[2:-2])
                run.bold = True
            elif parte.startswith('*') and parte.endswith('*'):
                run = p.add_run(parte[1:-1])
                run.italic = True
            elif parte.startswith('`') and parte.endswith('`'):
                run = p.add_run(parte[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            else:
                p.add_run(parte)
    
    return p

def adicionar_tabela_markdown(doc, linhas):
    """Converte tabela Markdown para tabela Word"""
    # Filtrar linhas vazias e separadores
    linhas_filtradas = [l for l in linhas if l.strip() and not re.match(r'^\|[\s:-]+\|$', l)]
    
    if len(linhas_filtradas) < 2:
        return
    
    # Extrair cabeçalho e dados
    cabecalho = [c.strip() for c in linhas_filtradas[0].split('|')[1:-1]]
    dados = []
    for linha in linhas_filtradas[1:]:
        cols = [c.strip() for c in linha.split('|')[1:-1]]
        if cols:
            dados.append(cols)
    
    # Criar tabela
    table = doc.add_table(rows=1 + len(dados), cols=len(cabecalho))
    table.style = 'Light Grid Accent 1'
    
    # Preencher cabeçalho
    for i, texto in enumerate(cabecalho):
        cell = table.rows[0].cells[i]
        cell.text = texto
        cell.paragraphs[0].runs[0].bold = True
    
    # Preencher dados
    for i, linha in enumerate(dados):
        for j, texto in enumerate(linha):
            if j < len(table.rows[i + 1].cells):
                table.rows[i + 1].cells[j].text = texto
    
    return table

def processar_markdown(doc, conteudo):
    """Processa conteúdo Markdown e converte para Word"""
    linhas = conteudo.split('\n')
    i = 0
    tabela_buffer = []
    bloco_codigo = []
    em_codigo = False
    
    while i < len(linhas):
        linha = linhas[i]
        
        # Bloco de código
        if linha.strip().startswith('```'):
            if em_codigo:
                # Finalizar bloco de código
                for cod_linha in bloco_codigo:
                    adicionar_paragrafo(doc, cod_linha, codigo=True)
                bloco_codigo = []
                em_codigo = False
            else:
                # Iniciar bloco de código
                em_codigo = True
            i += 1
            continue
        
        if em_codigo:
            bloco_codigo.append(linha)
            i += 1
            continue
        
        # Títulos
        if linha.startswith('#'):
            match = re.match(r'^(#{1,6})\s+(.+)', linha)
            if match:
                nivel = len(match.group(1))
                texto = match.group(2)
                adicionar_titulo(doc, texto, nivel)
                i += 1
                continue
        
        # Tabelas
        if '|' in linha and linha.strip().startswith('|'):
            tabela_buffer.append(linha)
            i += 1
            # Continuar coletando linhas da tabela
            while i < len(linhas) and '|' in linhas[i]:
                tabela_buffer.append(linhas[i])
                i += 1
            # Processar tabela
            adicionar_tabela_markdown(doc, tabela_buffer)
            tabela_buffer = []
            continue
        
        # Linha horizontal
        if linha.strip() in ['---', '***', '___']:
            doc.add_paragraph('_' * 50)
            i += 1
            continue
        
        # Lista
        if linha.strip().startswith(('-', '*', '+')):
            texto = linha.strip()[1:].strip()
            p = doc.add_paragraph(texto, style='List Bullet')
            i += 1
            continue
        
        # Linha vazia
        if not linha.strip():
            i += 1
            continue
        
        # Parágrafo normal
        adicionar_paragrafo(doc, linha)
        i += 1

def converter():
    """Função principal de conversão"""
    print("📄 Lendo arquivo Markdown...")
    
    if not MARKDOWN_FILE.exists():
        print(f"❌ Arquivo não encontrado: {MARKDOWN_FILE}")
        return False
    
    with open(MARKDOWN_FILE, 'r', encoding='utf-8') as f:
        conteudo = f.read()
    
    print("📝 Criando documento Word...")
    doc = criar_documento()
    
    print("🔄 Processando conteúdo...")
    processar_markdown(doc, conteudo)
    
    print("💾 Salvando arquivo...")
    doc.save(str(OUTPUT_FILE))
    
    size_kb = OUTPUT_FILE.stat().st_size / 1024
    print(f"✅ Conversão concluída!")
    print(f"📁 Arquivo: {OUTPUT_FILE}")
    print(f"📊 Tamanho: {size_kb:.1f} KB")
    
    return True

if __name__ == "__main__":
    print("=" * 60)
    print("  CONVERSOR MARKDOWN → DOCX (python-docx)")
    print("=" * 60)
    print()
    
    try:
        converter()
        print("\n🎉 Processo concluído!")
    except ImportError:
        print("❌ Erro: Biblioteca 'python-docx' não encontrada")
        print("\nInstalação:")
        print("  pip install python-docx")
    except Exception as e:
        print(f"❌ Erro: {e}")
        import traceback
        traceback.print_exc()
